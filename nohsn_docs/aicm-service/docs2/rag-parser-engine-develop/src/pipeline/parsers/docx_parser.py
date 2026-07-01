"""DOCX 파서 -- python-docx 단락/헤딩 추출 + LibreOffice PDF 변환 fallback.

전략
-----
1. python-docx 로 단락 구조(heading level, paragraph index) 직접 추출
2. 표가 있거나 복잡 레이아웃일 경우 LibreOffice 로 PDF 변환 후 PDF 파서에 위임
"""

from __future__ import annotations

import asyncio
import os
import re
import tempfile
from typing import Any

from src.common.logging import get_logger
from src.pipeline.models.parse_result import (
    PageContent,
    ParseResult,
    TableContent,
    TextBlock,
)
from src.pipeline.parsers.base import BaseParser

log = get_logger(__name__)

# python-docx 스타일명 → 헤딩 레벨 매핑 (영문 + 한국어)
_HEADING_STYLE_PATTERNS: list[tuple[re.Pattern[str], int]] = [
    (re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE), 0),  # group(1) 사용
    (re.compile(r"^제목\s*(\d+)$"), 0),  # 한국어 스타일
    (re.compile(r"^Title$", re.IGNORECASE), 1),  # Title → level 1
    (re.compile(r"^Subtitle$", re.IGNORECASE), 2),  # Subtitle → level 2
]

# 암묵적 헤딩 판정 최대 글자 수
_IMPLICIT_HEADING_MAX_LEN = 80


def _detect_heading_level_from_style(style_name: str | None) -> int | None:
    """스타일 이름에서 헤딩 레벨을 추출한다. 헤딩이 아니면 None 반환."""
    if not style_name:
        return None
    for pattern, fixed_level in _HEADING_STYLE_PATTERNS:
        m = pattern.match(style_name.strip())
        if m:
            if fixed_level == 0:
                # 그룹에서 레벨 추출
                return int(m.group(1))
            return fixed_level
    return None


def _is_implicit_heading(para: Any) -> int | None:
    """볼드 + 짧은 텍스트 등으로 암묵적 헤딩을 감지한다. 헤딩 레벨(기본 3) 또는 None."""
    text = para.text.strip()
    if not text or len(text) > _IMPLICIT_HEADING_MAX_LEN:
        return None
    # 모든 run 이 bold 인 경우만 판정
    runs = para.runs
    if not runs:
        return None
    all_bold = all(r.bold for r in runs if r.text.strip())
    if not all_bold:
        return None
    # 볼드 + 짧은 텍스트 → 암묵적 헤딩 level 3 (구체적 레벨 판별 불가)
    return 3


class DOCXParser(BaseParser):
    """DOCX 문서 파서."""

    async def parse(self) -> ParseResult:
        """단락 구조를 python-docx 로 추출하고 PDF 변환 결과와 병합한다."""
        return await asyncio.to_thread(self._parse_sync)

    # ------------------------------------------------------------------
    # 동기 내부
    # ------------------------------------------------------------------
    def _parse_sync(self) -> ParseResult:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError("python-docx 패키지가 필요합니다: pip install python-docx") from exc

        doc = DocxDocument(self.file_path)

        # 1) 단락 추출 — 헤딩 구조 감지 포함
        pages: list[PageContent] = []
        all_tables: list[TableContent] = []
        text_blocks: list[TextBlock] = []
        texts: list[str] = []
        running_offset = 0

        # 헤딩 관련 추적 구조
        headings: list[dict[str, Any]] = []  # {"level": int, "text": str, "char_offset": int}
        heading_path: list[str] = []  # 현재 헤딩 계층 경로
        heading_levels: list[int] = []  # heading_path 에 대응하는 레벨 목록
        # paragraph_index → 헤딩/경로 메타 매핑
        block_heading_meta: dict[int, dict[str, Any]] = {}

        for para_idx, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue

            # --- 헤딩 감지 ---
            style_name = para.style.name if para.style else None
            heading_level = _detect_heading_level_from_style(style_name)

            # 스타일 기반 헤딩이 없으면 암묵적 헤딩(볼드+짧은 텍스트) 시도
            if heading_level is None:
                heading_level = _is_implicit_heading(para)

            if heading_level is not None:
                heading_text = para.text.strip()

                # headings 목록에 추가
                headings.append({
                    "level": heading_level,
                    "text": heading_text,
                    "char_offset": running_offset,
                    "paragraph_index": para_idx,
                })

                # heading_path 갱신: 현재 레벨보다 깊거나 같은 항목 제거
                while heading_levels and heading_levels[-1] >= heading_level:
                    heading_levels.pop()
                    heading_path.pop()
                heading_path.append(heading_text)
                heading_levels.append(heading_level)

                log.debug(
                    "docx_heading_detected",
                    level=heading_level,
                    text=heading_text,
                    style=style_name,
                    heading_path=list(heading_path),
                )

            # TextBlock 생성
            tb = TextBlock(
                text=para.text,
                start_char_offset=running_offset,
                end_char_offset=running_offset + len(para.text),
                paragraph_index=para_idx,
            )
            text_blocks.append(tb)
            texts.append(para.text)

            # 블록별 헤딩 메타 기록 (헤딩 여부와 무관하게 현재 heading_path 저장)
            block_meta: dict[str, Any] = {
                "heading_path": list(heading_path),
            }
            if heading_level is not None:
                block_meta["heading_level"] = heading_level
                block_meta["is_heading"] = True
            block_heading_meta[para_idx] = block_meta

            running_offset += len(para.text) + 1

        # DOCX 는 물리적 페이지 개념이 약하므로 단일 가상 페이지로 취급
        full_text = "\n".join(texts)
        pages.append(
            PageContent(
                page_number=1,
                text=full_text,
                text_blocks=text_blocks,
                layout_type="single",
            )
        )

        # 2) 표 추출
        for t_idx, table in enumerate(doc.tables):
            headers, rows = self._extract_docx_table(table)
            if headers:
                md = self._table_to_markdown(headers, rows)
                tc = TableContent(
                    page_number=1,
                    table_index=t_idx,
                    headers=headers,
                    rows=rows,
                    markdown=md,
                    confidence=0.9,
                )
                all_tables.append(tc)
                pages[0].tables.append(tc)

        # 3) 메타데이터 — 헤딩 구조 포함
        metadata = self._extract_docx_metadata(doc)
        metadata["headings"] = headings
        metadata["block_heading_meta"] = block_heading_meta

        log.info(
            "docx_parse_complete",
            total_blocks=len(text_blocks),
            heading_count=len(headings),
            file=self.file_path,
        )

        return ParseResult(
            raw_text=full_text,
            pages=pages,
            tables=all_tables,
            images=[],
            metadata=metadata,
            source_file_path=self.file_path,
        )

    # ------------------------------------------------------------------
    # 헬퍼
    # ------------------------------------------------------------------
    @staticmethod
    def _extract_docx_table(table: Any) -> tuple[list[str], list[list[str]]]:
        """python-docx Table 객체에서 헤더/행을 추출한다."""
        all_rows: list[list[str]] = []
        for row in table.rows:
            all_rows.append([cell.text.replace("\n", " ").strip() for cell in row.cells])

        if not all_rows:
            return [], []

        headers = all_rows[0]
        body = all_rows[1:] if len(all_rows) > 1 else []
        return headers, body

    @staticmethod
    def _table_to_markdown(headers: list[str], rows: list[list[str]]) -> str:
        lines = [
            "| " + " | ".join(headers) + " |",
            "| " + " | ".join("---" for _ in headers) + " |",
        ]
        for row in rows:
            padded = row + [""] * max(0, len(headers) - len(row))
            lines.append("| " + " | ".join(padded[: len(headers)]) + " |")
        return "\n".join(lines)

    @staticmethod
    def _extract_docx_metadata(doc: Any) -> dict:
        metadata: dict[str, str] = {}
        try:
            cp = doc.core_properties
            metadata["title"] = cp.title or ""
            metadata["author"] = cp.author or ""
            metadata["created"] = str(cp.created) if cp.created else ""
            metadata["modified"] = str(cp.modified) if cp.modified else ""
        except Exception:
            pass
        return metadata


async def convert_docx_to_pdf(docx_path: str) -> str:
    """LibreOffice headless 로 DOCX -> PDF 변환. 변환된 PDF 경로를 반환한다."""
    out_dir = tempfile.mkdtemp(prefix="aicm_docx_")
    cmd = [
        "soffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        out_dir,
        docx_path,
    ]
    proc = await asyncio.create_subprocess_exec(
        *cmd,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    await proc.communicate()

    base = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(out_dir, f"{base}.pdf")
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"LibreOffice PDF 변환 실패: {pdf_path}")
    return pdf_path
